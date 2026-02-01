"""Generate production notes documents for recording."""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime


@dataclass
class ProductionNotes:
    """Structured production notes for a video recording session."""
    title: str
    duration_estimate: int  # minutes
    pre_recording_checklist: List[str]
    timing_summary: List[Dict[str, Any]]
    cue_legend: Dict[str, str]
    cell_alignment_table: List[Dict[str, Any]]
    visual_assets: List[str]


class ProductionNotesGenerator:
    """Generate production notes documents for recording."""

    CUE_LEGEND = {
        "[SHOW SLIDE: ...]": "Display the specified static slide image",
        "[SWITCH TO: Jupyter Notebook]": "Transition to the Jupyter notebook demo",
        "[RUN CELL]": "Execute the current notebook cell",
        "[PAUSE]": "Brief pause (1-2 seconds) for viewer absorption",
        "[STAY ON: ...]": "Keep current visual while continuing narration",
    }

    PRE_RECORDING_CHECKLIST = [
        "Open Jupyter Notebook in browser (zoom to 120% for readability)",
        "Clear all notebook outputs: Kernel > Restart & Clear Output",
        "Load sample dataset file into working directory",
        "Have slides ready in separate folder for quick switching",
        "Test screen recording captures full 1920x1080",
        "Disable notifications and popups",
        "Close unnecessary browser tabs and bookmarks bar",
    ]

    def generate(
        self,
        script,
        project_dir: Path,
        fmt: str = "docx",
    ) -> Path:
        """Generate production notes document.

        Args:
            script: An ImportedScript instance.
            project_dir: Project directory for locating assets.
            fmt: Output format — 'docx' or 'md'.

        Returns:
            Path to the generated file.
        """
        notes = self._build_notes(script, project_dir)

        if fmt == "docx":
            return self._write_docx(notes, project_dir / "production_notes.docx")
        return self._write_markdown(notes, project_dir / "production_notes.md")

    def _build_notes(self, script, project_dir: Path) -> ProductionNotes:
        """Build production notes from an ImportedScript."""
        timing = []
        for section in script.sections:
            word_count = len(section["text"].split())
            duration = max(30, int(word_count / 2.5))  # ~150 WPM
            timing.append({
                "segment": section["type"],
                "duration": f"{duration} seconds",
                "visual": ", ".join(section.get("visual_cues", [])[:2]) or "Talking head / Notebook",
                "narration_preview": section["text"][:100] + "...",
            })

        slides_dir = project_dir / "slides" / "png"
        visual_assets = (
            [f.name for f in sorted(slides_dir.glob("*.png"))]
            if slides_dir.exists()
            else []
        )

        return ProductionNotes(
            title=script.title,
            duration_estimate=script.duration_estimate,
            pre_recording_checklist=self.PRE_RECORDING_CHECKLIST,
            timing_summary=timing,
            cue_legend=self.CUE_LEGEND,
            cell_alignment_table=self._build_cell_alignment(script),
            visual_assets=visual_assets,
        )

    def _build_cell_alignment(self, script) -> List[Dict[str, Any]]:
        """Build cell-to-script alignment table."""
        return [
            {
                "cell": i,
                "section": block.get("section", ""),
                "code": block["code"][:50] + "...",
                "language": block.get("language", "python"),
            }
            for i, block in enumerate(script.code_blocks, 1)
        ]

    def _write_docx(self, notes: ProductionNotes, output_path: Path) -> Path:
        """Write production notes to Word document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # Fall back to markdown if python-docx not available
            md_path = output_path.with_suffix(".md")
            return self._write_markdown(notes, md_path)

        doc = Document()

        # Title
        title_para = doc.add_heading(f"Production Notes: {notes.title}", level=1)

        # Duration
        doc.add_paragraph(
            f"Estimated Duration: {notes.duration_estimate} minutes"
        )
        doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )

        # Pre-Recording Checklist
        doc.add_heading("Pre-Recording Checklist", level=2)
        for item in notes.pre_recording_checklist:
            p = doc.add_paragraph(item, style="List Bullet")

        # Timing Summary
        doc.add_heading("Timing Summary", level=2)
        if notes.timing_summary:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            header_cells[0].text = "Segment"
            header_cells[1].text = "Duration"
            header_cells[2].text = "Visual"
            header_cells[3].text = "Narration Preview"

            for entry in notes.timing_summary:
                row = table.add_row().cells
                row[0].text = entry["segment"]
                row[1].text = entry["duration"]
                row[2].text = entry["visual"]
                row[3].text = entry["narration_preview"]

        # Cue Legend
        doc.add_heading("Cue Legend", level=2)
        cue_table = doc.add_table(rows=1, cols=2)
        cue_table.style = "Table Grid"
        cue_table.rows[0].cells[0].text = "Cue"
        cue_table.rows[0].cells[1].text = "Meaning"
        for cue, meaning in notes.cue_legend.items():
            row = cue_table.add_row().cells
            row[0].text = cue
            row[1].text = meaning

        # Cell Alignment
        if notes.cell_alignment_table:
            doc.add_heading("Notebook Cell Alignment", level=2)
            cell_table = doc.add_table(rows=1, cols=4)
            cell_table.style = "Table Grid"
            header = cell_table.rows[0].cells
            header[0].text = "Cell"
            header[1].text = "Section"
            header[2].text = "Code Preview"
            header[3].text = "Language"

            for entry in notes.cell_alignment_table:
                row = cell_table.add_row().cells
                row[0].text = str(entry["cell"])
                row[1].text = entry["section"]
                row[2].text = entry["code"]
                row[3].text = entry["language"]

        # Visual Assets
        if notes.visual_assets:
            doc.add_heading("Visual Assets", level=2)
            for asset in notes.visual_assets:
                doc.add_paragraph(asset, style="List Bullet")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _write_markdown(self, notes: ProductionNotes, output_path: Path) -> Path:
        """Write production notes as markdown."""
        lines = [
            f"# Production Notes: {notes.title}\n",
            f"**Estimated Duration:** {notes.duration_estimate} minutes  ",
            f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n",
            "## Pre-Recording Checklist\n",
        ]

        for item in notes.pre_recording_checklist:
            lines.append(f"- [ ] {item}")

        lines.append("\n## Timing Summary\n")
        lines.append("| Segment | Duration | Visual | Narration Preview |")
        lines.append("|---------|----------|--------|-------------------|")
        for entry in notes.timing_summary:
            lines.append(
                f"| {entry['segment']} | {entry['duration']} "
                f"| {entry['visual']} | {entry['narration_preview']} |"
            )

        lines.append("\n## Cue Legend\n")
        lines.append("| Cue | Meaning |")
        lines.append("|-----|---------|")
        for cue, meaning in notes.cue_legend.items():
            lines.append(f"| `{cue}` | {meaning} |")

        if notes.cell_alignment_table:
            lines.append("\n## Notebook Cell Alignment\n")
            lines.append("| Cell | Section | Code Preview | Language |")
            lines.append("|------|---------|--------------|----------|")
            for entry in notes.cell_alignment_table:
                lines.append(
                    f"| {entry['cell']} | {entry['section']} "
                    f"| `{entry['code']}` | {entry['language']} |"
                )

        if notes.visual_assets:
            lines.append("\n## Visual Assets\n")
            for asset in notes.visual_assets:
                lines.append(f"- {asset}")

        content = "\n".join(lines) + "\n"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")
        return output_path
